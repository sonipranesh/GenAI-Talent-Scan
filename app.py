import base64
import os
from turtle import pd
import fitz
import pandas as pandas
import streamlit as st
import docx2txt
import re
import gemini_api as genai 
#from pyresparser import ResumeParser
from io import BytesIO
from datetime import datetime
from threading import Thread 
from queue import Queue 
from bokeh.models import ColumnDataSource, CustomJS
from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_PARAGRAPH_ALIGNMENT
import json


#Below 3 lines of code are used to declare the Gemini Engine
GOOGLE_API_KEY = "YOUR API KEY"
genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-pro')


#This function is used to extract text from the PDF.
def extract_text_from_pdf(pdf_path):
    try:

        with fitz.open(pdf_path) as pdf_document:

            text = "".join([page.get_text() for page in pdf_document])
        
        return text

    except Exception as e:

        #st.error(f"Error processing {pdf_path}: {e}")

        return None

#This function is used to extract text from the doc
def extract_text_from_docx(docx_path):
    try:
        text = docx2txt.process(docx_path)
        #doc= Document(docx_path)
        #text = ""
        #for paragraph in doc.paragraphs:
        #    text += paragraph.text + "\n"
        return text

    except Exception as e:
        #st.error(f"Error processing {pdf_path}: {e}")
        return None


def extract_text_from_doc(docx_path):
    try:
        doc= Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
           text += paragraph.text + "\n"
        return text

    except Exception as e:
        #st.error(f"Error processing {pdf_path}: {e}")
        return None


#This function is used to store the resumes in the form of dictionary from the provided folder path.
def load_resumes_as_dict(folder_path):
    resumes_dict = {}

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path,file_name)
        if file_name.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)

        elif file_name.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        
        #elif file_name.endswith('.doc'):
        #     text = extract_text_from_doc(file_path)

        else:
            continue

        if text:
            resumes_dict[file_name] =(file_path,text)
    
    return resumes_dict


#This is the batch-processing version of the above function.
def load_resumes_as_dict2(folder_path, batch_files):
    resumes_dict2 ={}

    for file_name in batch_files:
        file_path = os.path.join(folder_path,file_name)
        if file_name.endswith('.pdf'):
            text= extract_text_from_pdf(file_path)
        
        elif file_name.endswith('.docx'):
            text= extract_text_from_docx(file_path)
        
        else:
            continue

        if text:
            resumes_dict2[file_name] = (file_path, text)

    return resumes_dict2


data = []


#This function generates the resume summary
def generate_resume_summary(resumes_dict,optional_skills,datascience_skills,resume_folder):
    for file_name,(file_path,text_content) in resumes_dict.items():

        #parsed_resume = ResumeParser(file_path).get_extracted_data()

        #skills_used = ', '.join(parsed_resume.get('skills', 'Not Provided'))
        #st.write(skills_used)
        #prompt05=f"""
        #         Act as a HR and your job is to check the resumes for data science.
        #
        #         Provided the answer for the following question:
        #         1. Please list down the core skills, technologies, libraries caandidate should have in their resume to get selected for Data Science job profile?
        #
        #         """
        #new = model.generate_content(prompt05,generation_config=genai.types.GenerationConfig(
        #      candidate_count=1,
        #      temperature=0.3)
        #                                    )
        #st.write(new.text)

        prompt4= f"""
                   Your task is to analyze the resume content with a focus on skills related to Data Science, Gen AI, NLP, or Machine Learning.
                   Extracting 'Duration' is critical part of this task so use your understanding to extract dates for each project.

                   Provide a concise summary into two tables, first for work experience/professional experience and second for projects structured as follows:

                   1. Refer projects and work experiences related sections of the resumes. Only extract projects/work experience/ professional experience related to data science, machine learning, genai ONLY if he/she 'developed' or 'build' something in the project. Highlight the duration for each such projects that has a specified period (year) and infer the domain if not explicitly mentioned. Domain should be industry related only like 'Life Sciences', 'Healthcare'. Please avoid using generic domains like 'Computer Vision' or 'NLP'.

                   For each project and work experience/professional experience:
                        - Extract only the skills mentioned for that particular project aand the role played. 
                        - Extract only the GenAI related skills mentioned for that particular project ONLY if these skills were utilized in AI, GenAI projects or have used skills like OpenAI, OpenAI API, GPT, GPT API, Gemini, Gemini API, Langchain, prompt engineering, LLM models, etc in the project. Use your understanding of GenAI skills.
                        - List only the Data Science related skills present in the relevant project/work experience/ professional experience ONLY if these skills are closely related to any of the following skills, along with use your understanding of data science skills: {datascience_skills}.
                        - For Preferred Requirement, extract 'GCP' or 'AWS' or 'Azure' ONLY if the candidate has demonstrable experience with cloud services unrelated to database management. 
                        - For Additional skills, check if the candidate has mentioned or used skills {optional_skills} in the resume, if yes then provide that skill along with the experience in years of the candidate for that skills else return 'None'. 
                        - For Alternate Recommendationm refer onlky {optional_skills} skills, experience of the candidate in this skills from the list, where and how candidate use this skills from the list; based on this recommend the candidate for this skills alongwith short explaination.
                        - Extract only 'Python' or 'R' Programming languages mentioned in the resume.

                    Use your understanding, do not extrtact skills mentioned in the technical skills in 'Skills utilized', 'GenAI related', 'Data Science' related, strictly extracrt only the skills which are mentioned under project description or responsibilities for that particular project related to only genai and data science. 
                    Please focus only extracting the skills from the options I have provided for Preferred Requirement and Additional Skills.
                    Please focus to extract experience in years for Additional skills. If {optional_skills} is empty then provide Additional Skills and Alternate recommedation as 'None'. 
                    And it is must to ensure the extraction of project durations that is 'start date' and 'end date' mentioned by the candidate, which will be mentioned after or before the title of the project or under one duration multiple project would be mentioned then consider same time period for each project. Consider 'Present' as January 2024. 
                    Please focus to strictly extract project or work experience duratrion only if candidate has used Data Science or GenAI related skills in the project else return None. 
                    For projects where periods are not mentioned, indicate the duration as 'None'. 
                    For role, write only data science related, data analysis related or generative AI related job roles. For example, job roles related to data science are: Data Scientist, Machine Learning Engineer, Generative AI Engineer, etc. For other roles write "Not Relevant" in role column only. 
                    For Role Experience(in years), Calculate the duration of the intercval mentioned, consider present as Feb 2024.

                    For the presence of 'Git' or 'Github' or Github link anywhere in the resume, provide only 'Yes' or 'No'. 
                    Always provide 'Name of the candidate' from the resume. When you are not able to extract "Name of the Candidate" from resume, use {file_name}.
                    Always provide both 'Name of the Candidate' and 'GitHub/link' below the project table only, pleaase don't include them into tables. 

                    Please focus candidate may have written 'Work Experience' in separate section like 'Career', 'Carrer progression', 'Organisational Experience' or 'Experience', use your understanding and include it in 'work experience'. 
                    Avoid including educational experience in 'work experience'. 

                    Always ensure to provide 'Programming Languages' and 'Name of Candidate'. 

                    Exclude the provided example from the analysis.NameError

                    For example:

                        Project: Designing IoT face Recognition Robot
                        Role: Data Scientist
                        Role Experience (in years): 1.3
                        Details: Utilizex skills in facial recognition, creating a training database, and recording attendance and temperature. 
                        Duration: August'2015-December'2016
                        Skills Utilized: Python
                        GenAI related: None
                        Data Science related: CNN, TensorFlow, Random Forest, Machine Learning
                        Preferred Requirement: AWS
                        Aadditional Skills: None
                        Alternate Recommendation: None
                        Domain: Robotics
                        Programming Languages: Python
                        Additional Information: "Name of Candidate: Pranesh Soni",Github/Link: Yes"

                    Provide the analysis in a structured tabular format as outlined above. Focus solely on the relevant information requested and avoid addtional detaails. 

                    Here is the resume content: {text_content}.

                    """
        
        try:

            # response = openai.completion.create(
            #
            # engine= "text-davinci-003",
            # 
            # prompt = prompt4,
            #
            # max_tokens = 500,
            # 
            # temperature= 0
            # 
            # )
            #
            # response_text1 = response.choice[0].text.strip()
            #st.write(response_text1)

            response_text = model.generate_content(prompt4, generation_config= genai.types.GenerationConfig(
                candidate_count= 1,
                temperature_count= 0.3)
                                                  )
            
            #st.write(response_text.text)

            excluded_terms = ['Teaching', 'Professor', 'Researcher', 'Assistant Professor', 'Trainer',
                             'RESEARCH SCHOLAR', 'intern', 'Intern', 'Quality Analyst']
            rows= response_text.text.strip().split('\n')

            def include_row(row):
                for term in excluded_terms:
                    if term.lower() in row.lower():
                        return False
                return True
            
            filtered_rows = [row for row in rows if include_row(row)]
            filtered_output = '\n'.join(filtered_rows)
            #st.write(filtered_output)

            # prompt01 = f"""
            #             Your task was to analyze the resume content with a focus on skills related to Data Science, Gen AI, AI, NLP or Machine Learning.
            #             Resume content and response generated by you is given below.
            #
            #             Check the answer for the below questions and provide updated response by making necessary changes in that table.
            #             1. What are the data science skills and genai skills?
            #             2. Are you sure skills extracted by you in the respective columns are correct? 
            #             3. Are you sure candidate has used that in the project?
            #
            #
            #             resume_content = {text_content}  
            #             your_response={filtered_output}
            #             """
            # Updated= model.generate_content(prompt01, generation_config=genai.types.GenerationConfig(
            #       candidate_count=1,
            #       temperature= 0.3)
            #       )
            #  
            # st.write(Updated.text)

            prompt2 = f"""
                     Your task is to use both 'Work Experience' and 'Projects' tables for extracting all the below information. 
                     
                     Please provide a concise summary of the candidate's  projects, calculate the total project and work experience duration in years, and extract the skills used across all projects and work experience. 
                     Please focus to provide a response in plain text format only without any formatting such as bold, italics, or special characters.
                     Always provide responses in comma sepaarated manner, Don't provide responses in bullets.

                     For the summary, consider using this format as a reference: "The candidate has worked on several projects, including Project 1, a virtual project focusing on ..., and Project 2, an application developed in Swift that..." 

                     1. Summary: (Provide the projects done by the candidate)

                     2. Duration: (Provide duration in years like 1 year, 2 year. Duration for 'Technical lead', 'Data Scientist', 'NLP developer' and other related fields from 'Work Experience' and add them up. Consider only one year if there are overlapping years. Don't take duration from 'projects'. )

                     3. Skills: (Extract thee skills and programming languages utilized by the candidate acrooss all projects and work experiences)

                    4. GenAI Related: (Extract the 'GenAI related' skills utilized by the caandidate across all projects and work experiences)

                    5. Data Science Related: (Extract the 'Data Science relaated' skills utilized by the candidate across all projects and work experineces)

                    6. Preferred Requirement: (Extract the 'Preferred Requirement' utilized by the candidate across all projects)

                    7. GitHub: (Provide 'Yes' or 'No' only from 'Github/link' column)

                    8. Additional Skills: (Extract the only 'Additional Skills' along with years utilized by the candidate from 'Additiona Skills' column only)

                    9. Alternate Recommendation: (Extract 'Alternate Recommendation' ffrom 'Alternate Recommendation' column)

                    10. Programming Languages: (Extract nlt the Programming Languages 'Python' or 'R' from skillss by the canadidate across all projects and work experiences or from 'Programming Languaages' columnn )

                    11. Domain: (Extract the 'Domain' utilized by the candidate across all projects)

                    12. Name: (Extract the 'Name of the candidate')

                    13. Role: (Extract all 'Role' from 'Role' column)

                    14. Experience: (Calculate the total duration from 'Role Experience(in years)' column, avoid considering same duration)

                    Please focus on providing the Summary always. Always ensure that extracted information aligns accurately with the respective fields.
                    Only provide above 14 information, don't provide extra information.


                    Always verify your answer by checking data science and genai skills are correctly extracted or not before submitting.
                    Always focus, If both data science and genai related skills are None then provide None for all other columns also except Name.

                    summary: {filtered_output}

                    """

            # response= openai.completion.create(
            # 
            # engine = "text-davinci-003",
            #
            # prompt = prompt2,
            #
            # max_tokens = 500,
            #
            # temperature = 0.1
            #
            # )     
            #
            # response_text2= response.choice[0].text.strip()
            #st.write(response_text2)

            # 4.
            # GenAI
            # Related: ( Extract
            #           the 'GenAI' or 'GenAI related skills' utilized by the candidate from 'GenAI Related' column only)
            # 5. 
            # Data 
            # Science
            # Related: (Extract
            #           the 'Data Science' or 'Data Science related skills' utilized by the candidate from 'Data Science Related' column only)

            response_text1= model.generate_content(prompt2, generation_config = genai.types.GenerationConfig(
                candidate_count=1,
                top_p= 0.6,
                top_k=5,
                temperature=0)
                                        )

            # st.write(response_text1.text)

            # prompt02 = f"""
            #             Your job was to extract skills mentioned in actual data to respective columns.
            #
            #             Check the answer for the below questions and provide updated response by making necessary changes in your response only not in actual data.
            #             1. are the skills extracted in respective columns in your response match with columns in the actual data?
            #             2. does data science and genai related skills are extracted accurately in your response?
            #
            #
            #
            #
            #             Your response={response_text1}
            #             Actual data={filtered_output}
            #             """
            #
            # updated1 = model.generate_content(prompt02, generation_config=genai.types.GenerationConfig(
            #     candidate_count=1,
            #     top_p=0.6,
            #     top_k=5,
            #     temperature=0)
            #                                         )
            #
            # st.write(updated1.text)

            summary_pattern = r"\*{0,2}Summary:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Duration:\{0,2}|\n)"
            
            duration_pattern = r"\*{0,2}Duration:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Skills:\*{0,2}|\n)"
            
            skills_pattern = r"\*{0,2}Skills:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}GenAI:\*{0,2}|\n)"
            
            genai_pattern = r"\*{0,2}GenAI:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}DataScience:\*{0,2}|\n)"
            
            datascience_pattern = r"\*{0,2}DataScience:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Preferred:\*{0,2}|\n)"
            
            preferred_pattern = r"\*{0,2}Preferred:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}GitHub:\*{0,2}|\n)"
            
            github_pattern = r"\*{0,2}GitHub:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}AdditionalSkills:\*{0,2}|\n)"
            
            additionalskills_pattern = r"\*{0,2}AdditionalSkills:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}AlternateRecommendation:\*{0,2}|\n)"
            
            alternate_recommendation_pattern = r"\*{0,2}AlternateRecommendation:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}End:\*{0,2}|\n)"
            
            programming_languages_pattern = r"\*{0,2}ProgrammingLanguages:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Domain:\*{0,2}|\n)"
            
            domain_pattern = r"\*{0,2}Domain:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Name:\*{0,2}|\n)"
            
            name_pattern = r"\*{0,2}Name:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Role:\*{0,2}|\n)"
            
            role_pattern = r"\*{0,2}Role:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}Experience:\*{0,2}|\n)"
            
            experience_pattern = r"\*{0,2}Experience:\*{0,2}\s(.*?)(?=\d+\. |\*{0,2}End:\*{0,2}|\n)"

            # Extract text following each label using regular expressions

            extract_text = lambda pattern, response: re.search(pattern, response, re.DOTALL).group(
                
                1).strip() if re.search(pattern, response, re.DOTALL) else "Could not read"  
            
            summary= extract_text(summary_pattern, response_text1.text)

            duration = extract_text(duration_pattern, response_text1.text)

            skills = extract_text(skills_pattern, response_text1.text)

            GenAI = extract_text(genai_pattern, response_text1.text)
            #st.write(GenAI)

            DataScience = extract_text(datascience_pattern, response_text1.text)

            Preferred = extract_text(preferred_pattern, response_text1.text)

            gitHub = extract_text(github_pattern, response_text1.text)

            Additional_Skills = extract_text(additionalskills_pattern, response_text1.text)
            Additional_Skills = None

            alternate_recommendation = extract_text(alternate_recommendation_pattern, response_text1.text)
            alternaate_recommendation = None

            programming_languages = extract_text(programming_languages_pattern, response_text1.text)

            domain = extract_text(domain_pattern, response_text1.text)

            name1 = extract_text(name_pattern, response_text1.text)

            name = name1.capitalize()

            if str(name).lower() in ['none', 'not provided', 'not mentioned', 'could not read']:
                name = ' '.join(re.sub(r'\[.*?\]', '', os.path.splitext(file_name)[0]).replace('+', ' ').replace('_', ' ').replace('Naukri', ' ').split()[:2]).capitalize()
            else:
                name = name

            role = extract_text(role_pattern, response_text1.text)

            experience = extract_text(experience_pattern, response_text1.text)

            #Extracting recommendation

            def calculate_skill_score(skills_str):
                skill_count =0
                skill_count = skills_str.count(',') + 1

                score = 0 
                if skill_count == 1 and (
                    'None' in skills_str or 'Not' in skills_str or 'No' in skills_str or 'not' in skills_str or skills_str ==""):
                    skill_count =0 
                    score = 0

                
                elif skill_count >=1 and skill_count <=4:
                    score = 0.25 * skill_count
                else:
                    score = 1
                
                return [score, skill_count]
            
            def calculate_exp_score(duration):
                
                if duration == 0:
                    score = 0
                
                elif duration >= 1 and duration <= 2:
                    score = 0.25
                
                elif duration > 2 and duration <= 5:
                    score = 0.5 
                
                elif duration > 5 and duration <= 10:
                    score = 0.75
                
                else:
                    score = 1
                
                return score
            
            def calculate_total_score(g, d, c):

                total_score = g + d + c

                if total_score > 0.5:

                    recommendation = "Recommended"
                
                elif total_score == 0.5:

                    recommendation = "Border Line"
                
                else:

                    recommendation = "Not Recommended"
                
                return total_score, recommendation 
            
            def new_recommendation(genaiScore, datascienceScore, CloudScore, totalScore):

                if genaiScore == 0 and datascienceScore == 0:
                    comment = "No. Candidate doesn't have essential skills."

                elif (genaiScore >= datascienceScore) and totalScore > 0.5:
                    comment = "Yes. Recommended for GenAI profile."
                
                elif (genaiScore < datascienceScore) and totalScore > 0.5 and datascienceScore >= 0.5:
                    comment = "Yes. Recommended for Data Science profile."
                
                elif (genaiScore >= datascienceScore) and totalScore == 0.5:
                    comment = "Borderline. Recommended for GenAI profile."
                
                elif (genaiScore <= datascienceScore) and (totalScore >= 0.5 and totalScore <= 0.75) and datascienceScore > 0.5:
                    comment = "Borderline. Recommended for Data Science profile."
                
                else:
                    comment = "No. Candidate doesn't have essential skills."
                
                return comment

            #Calculating commas in genAI skills

            genAI_skills = calculate_skill_score(str(GenAI))
            #st.write(genAI_skills)
            #st.write(genAI_skills[1])
            #calculating commaas in data science skills

            datascience_skills = calculate_skill_score(str(DataScience))
            #st.write(datascience_skills)
            #st.write(datascience_skills[1])

            #Calculating commaas in cloud skills

            cloud_skills = calculate_skill_score(str(Preferred))
            #st.write(cloud_skills)

            #experience_score = calculate_exp_score(duration)
            #st.write(experience_score)

            total_score, recommendation = calculate_total_score(genAI_skills[0], datascience_skills[0], cloud_skills[0])
            #st.write(total_score)
            #st.write(recommendation)

            recommendation_from_score_text = new_recommendation(genAI_skills[0], datascience_skills[0], cloud_skills[0], total_score)

            
            if duration.lower() not in ['none', 'not mentioned', 'could not read', 'error in processing', 'not relevant']:
                duration_value = duration.split()[0].replace('+', '')
                if float(duration_value) < 1:
                    recommendation_from_score_text = "No. Candidate doesn't have essential skills."
                else:
                    pass
            else:
                pass

            if 'Yes' in recommendation_from_score_text:
                prompt99 = f"""
                       Consider you are an interviewer, your task is to generate potential questions to be asked from the candidate whose resume is {text_content}.

                        Generate a set of 10 comprehensive interview questions based on resume provided. Questions must be related to Data Science skills and/or GenAI skills for a candidate applying for a role either of Data Science and/or Generative AI. These questions should assess the candidate's understanding, experience, and technical skills across candidate's major skills:

                        1. Projects and Technical Aspects (5 Questions): Questions should delve into the candidate's practical experience and the technical challenges they've navigated in their projects. Generate typical technical questions as an interviewer based on projects/work experiences done by the candidates.

                        2. Python Specific Technical Skills (3 Questions): Ask conceptual programming questions to check candidate's programming ability and problem solving power in python.

                        3. Additional questions (2 questions): Ask questions based on other skills/cloud if mentioned in resume. for example, "How would you deploy machine learning model on AWS?"

                        Provide response in the tabular format. First row should contain only one column where name of the candidate should be mentioned.
                        Next rows should contain 2 columns , first column for questions with question number like Q1,Q2,etc, second column for score and it should be empty always.

                        Exclude the provided example from the analysis, Use the example only for the formatting.

                        For example:
                        "
                        | Name of the candidate |
                        | --- |
                        | DEBASIS SAMAL |
                        | Questions | Score |
                        | --- | --- |
                        | Q1.Explain the concept of Generative AI and its applications in the real world. |  |
                        | Q2.Describe your experience in developing the multilingual chatbot using Generative AI (OpenAI) on AWS SageMaker. How did you handle the challenges of multilingual processing and knowledge extraction from PDF documents? |  |
                        | --- | --- |
                        | ... |  |
                        "


                        Name of the candidate = {name}
                        GenAI skills = {GenAI}
                        Data Science skills = {DataScience}

                        """
                questions_text = model.generate_content(prompt99, generation_config=genai.types.GenerationConfig(
                    candidate_count=1,
                    temperature=0.3)
                )

                #st.write(question_text.text)
                #st.write(type(question_text.text))

                #Parse repsonse and extract data
                lines = questions_text.text.strip().split('\n')
                candidate_name = lines[2].strip().split("|")[1].strip()
                questions_and_answers = [line.strip() for line in line[4:]]

                # Create a new Document
                doc = Document()

                sections = doc.sections
                for section in sections:
                    section.tleft_margin = Pt(35) # Adjust left margin as needed
                    section.right_margin = Pt(0.5) # Adjust right margin as needed
                    section.top_margin = Pt(0.5) # Adjust top margin as needed
                    section.bottom_margin = Pt(0.5) # Adjust bottom margin as needed

                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                #doc.add_paragraph(f" Report Genearated At: {timestamp}", style='Heading 1')

                timestamp_paragraph = doc.add_paragraph(f"Report Generated At: {timestamp}")
                timestamp_run = timestamp_paragraph.runs[0]
                timestamp_run.bold = True
                timestamp_run.font.size = Pt(11) #Adjust the font size as needed

                # Add candidate naame
                doc.add_heading('Candidate Name:', level=1)
                doc.add_paragraph(candidate_name)

                # Add table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Questions'
                hdr_cells[1].text = 'Score'

                # Make column names bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                         for run in paragraph.runs:
                            run.bold = True

                # Populate Table
                for line in questions_and_answers:
                    parts = line.split('|')
                    if len(parts) > 1 and parts[1].strip() != '---': # Check if the line is not just '---'
                        question = parts[1].strip()
                        score = parts[2].strip() if len(parts) > 2 else '' # Handle the case where score is not provided
                        row_cells = table.add_row().cells
                        row_cells[0].text = question
                        row_cells[1].text = score

                        # Add space after each question 
                        row_cells[0].paragraphs[0].runs[-1].add_break()

                        # for cell in row_cells:
                        #    for paragraph in cell.paragraphs:
                        #    paragraph.space_after = Pt(10)

                # Adjust row height to fit content
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Set alignmnet to the left
                            paragraph.space_after = Pt(12) # Add space after each paragraph
                        cell.height = Pt(60) # Adjust row height as needed

                # Add a border around the entire document
                for paragraph in doc.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph.space_after = Pt(12)
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0,0,0) # Set font color to black
                        run.font.name = 'Cambria'
                        run.font.size = Pt(11)

                # Add space after the headings "Questions" and "Score"
                for cell in table.rows[0].cells:
                    cell.paragraphs[0].runs[-1].add_break()
                
                for cells in table.columns[1].cells:
                    cell.width = Pt(30) # Adjust column width as needed

                for cell in table.columns[0].cells:
                    cell.width = Pt(700)

                table.border_collapse = True

                new_folder = os.path.join(resume_folder, 'Generated_Docx')
                os.makedirs(new_folder, exist_ok=True)
                docx_filename = os.path.join(new_folder, f"{name}_reports.docx")
                doc.save(docx_filename)
                download_link = f"file:///{docx_filename}"
                docx_link = f'=HYPERLINK("{download_link}", "Open Report")'

            else:
                docx_link = ' '   

            prompt1= f"""
            Based on the following data,return 'Yes' or 'No' to indicate whether the profile is suitable for a data science-related position.

            Following are the conditions for positive recommendation:

                1.Candidate must have any of the 'Python' or 'R' in 'Skills' or 'Programming Languages'.

                2.Candidate must have either of 'GenAI Related' or 'Data Science Related' skills.

            Indicate 'Yes' only if the candidate satisfies two conditions above else 'No'.
            Indicate 'No' if only 'Machine Learning' is present in 'Data Science Related' skills and has 'None' in 'GenAI Related' skills.

            Also, provide a short reason for your recommendation.
            If 'GenAI Related' skills present then recommend candidate for GenAI profile.
            If only 'Data Science Related' skills present then recommend candidate for Data Science profile.
            Avoid using phrases like 'Answer:'.
            If Recommendation is 'No', provide reason as 'Candidate doesn't have essential skills'.

            Please focus exclude the provided example from the analysis.

            For example:
            Yes. Recommended for Data Science profile.
            Reason: Given the strong foundation and candidate demonstrates expertise in Data Science skills, the candidate is well-suited for positions related to Data Science.
            The candidate possesses proficiency in the Programming Languages, specifically in 'Python.'
            Notably, the candidate has cloud skills of AWS,Azure.

            Focus solely on the relevant information requested and avoid additional details.

            Summary: {response_text1.text}

            """

            # recommendation = openai.Completion.create(
            
            # engine - "text-davinci-003",

            # prompt = prompt1,

            # max_tokens = 500,

            # temperature = 0

            #  )                            
            
            # recommendation_text = recommendation.choice[0].text.strip()

            # recommendation_text = model.generate_content(prompt1, generation_config=genai.types.GenerationConfig(
            #    candidate_count=1,
            #    top_p=0.6,
            #    top_k=5,
            #   temperature=0)
            #  )

            # recommendation_text = process1(response_text, prompt1)

            #role,experience, datascience_skills[1], recommendation,

            pdf_link = f'{file_path}'
            data.append(file_name, name, summary, duration, skills, GenAI, DataScience, Preferred, gitHub,
                         programming_languages, domain, recommendation_from_score_text, Additional_Skills, 
                         alternate_recommendation, docx_link, pdf_link)
            
            # Print statements to check progress
            #st.write(data)
            print(f"Processed {file_name}")

        except Exception as e:
            
            # st.error(f"Error processing {file_name}: {e}")

            data.append(
                [file_name, 'Error in processing', 'Error in processing', 'Error in processing', 'Error in processing',
                 'Error in processing','Error in processing', 'Error in processing', 'Error in processing',
                 'Error in processing', 'Error in processing', 'Error in processing',
                 'Error in processing', 'Error in processing', 'Error in processing', 'Error in processing'])
    return data

def generate_questions(recommended_dict):
    questions = []
    for file_name, content in recommended_dict.items():

        prompt9 = f"""
                Generate a set of 15 comprehensive interview questions based on resume provided for a candidate applying for a role involving Data Science and/or Generative AI. These questions should assess the candidate's understanding, experience, and technical skills across three distinct areas:

                1. Fundamentals and Reasoning-Based Questions (First 5 Questions): Questions in this section should test the candidate's grasp of core concepts in data science and generative AI, including their ability to explain fundamental principles, methodologies, and theoretical aspects.Generate questions based on projects/work experiences done by the candidates.

                2. Projects and Technical Aspects (Next 5 Questions): This section should delve into the candidate's practical experience and the technical challenges they've navigated in their projects. Generate questions based on projects/work experiences done by the candidates.

                3. Python or R Specific Technical Skills (Last 5 Questions): Focus on assessing the candidate's proficiency in programming languages essential for data science and AI, such as Python or R. Generate questions based on projects/work experiences done by the candidates.

                Resume Content = {content}
                """

        try:
            questions_text = model.generate_content(prompt9, generation_config=genai.types.GenerationConfig(
                candidate_count=1,
                temperature=0.3)
                                                    )

            # st.write(questions_text.text)

            questions.append([file_name, questions_text.text])


        except Exception as e:

            # st.error(f"Error processing summary for {file_name}: {e}")

            questions.append([file_name], "Error in processing")

    return questions

#This function is rudimentary.
def rank_resume(data):
    prompt7 = f"""
        Analyze all the resumes given below for skills related to data science, genai,ai,ml.
        provide Confidence score out of 10 for each of the resume based on the GenAI related skills and Data Science related skills.

        resume data={data}
"""
    Rank = model.generate_content(prompt7, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        top_p=0.6,
        top_k=5,
        temperature=0))

    return Rank.text

#This function is rudimentaryt.
def resume_filter(resumes_dict, Optional_skills):
    fill = []
    for file_name, (file_path, text_content) in resumes_dict.items():
        # parsed_resume = ResumeParser(file_path).get_extracted_data()

        # skills_used = ', '.join(parsed_resume.get('skills', 'Not Provided'))
        # st.write(skills_used)
        # skills_used_list = [skill.strip().lower() for skill in skills_used.split(',')]
        # # st.write(skills_used_list)
        # optional_skills_lower = [skill.strip().lower() for skill in Optional_skills]
        # for skill in optional_skills_lower:
        #     if skill in skills_used_list:
        #         fill.append(file_name)

        prompt5 = f"""
                Please analyze the resume content for searching skills {Optional_skills} in the resume.

                Please check if candidate has mentioned or used skills {Optional_skills} in the resume, if yes then return {file_name} of the candidate, else return 'None'.
                if {Optional_skills} id empty then return 'None'.

                Focus solely on the relevant information requested and avoid additional details.

                Here is the resume content: {text_content}.

                """

        response = model.generate_content(prompt5, generation_config=genai.types.GenerationConfig(
            candidate_count=1,
            top_p=0.6,
            top_k=5,
            temperature=0)
                                          )
        fill.append(response.text)

    return fill

#This is a rudimentary function.
def Shortlisted_Resumes1(data):
    prompt3 = f"""
                Utilize the provided resume data containing 'File Name' and 'Recommendation' for multiple candidates.

                Total number of Resumes: (Extract the total count of PDF files (e.g Chandanverma.pdf) mentioned in 'File Name'.

                Please analyze the provided json format data, use your understanding and extract the following information:

                    - Number of Resumes suited for Data Science job profile: (Calculate the total count of 'File Name' where 'Recommendation' starts with word 'Yes'.)
                    - Number of Not suitable resumes: (Calculate the total count of 'File Name' where 'Recommendation' starts with word 'No')
                    - File names suitable for Data Science job profile: (Provide a list of all file names where 'Recommendation' starts with word 'Yes'.)

                Please ensure accuracy in counting, always verify your answer.
                Exclude the provided example from the analysis.

                For example:

                Total number of Resumes: 3
                Number of Resumes suited for Data Science job profile: 3
                Number of Not suitable resumes: 0
                File names suitable for Data Science job profile: Chandanverma.pdf, Chethan N.pdf, Yashaswini Kulkarni.pdf

                Provide the analysis in a structured format as outlined above.
                Please focus solely on extracting the requested information.

                Here is the provided data: {data}

                """

    response_text = model.generate_content(prompt3, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        temperature=0.3)
                                           )

    return response_text.text

#This is a rudimentary function.
def Shortlisted_Resumes(data):
    prompt3 = f"""
            Utilize the provided resume data containing details of multiple candidates in the list inside the list.

            Total number of Resumes: (Extract the total count of PDF files (e.g Chandanverma.pdf) mentioned in the data.Please focus for Counting 'Total number of resumes', avoid considering index of the list as it starts from 0.)

            Please analyze the provided data, use your understanding and extract the following information:

                - Number of Resumes suited for Data Science job profile: (Extract the total count of resumes where the last element in each resume's list starts with word 'Yes.'.)
                - Number of Not suitable resumes: (Extract the total count of resumes where the last element in each resume's list starts with word 'No.' or 'Unknown'.)
                - File names suitable for Data Science job profile: (Provide a list of file names for resumes where the last element in each resume's list starts with word 'Yes.'.)

            Exclude the provided example from the analysis.

            For example:

            Total number of Resumes: 3
            Number of Resumes suited for Data Science job profile: 3
            Number of Not suitable resumes: 0
            File names suitable for Data Science job profile: Chandanverma.pdf, Chethan N.pdf, Yashaswini Kulkarni.pdf

            Please focus to go through each of the resume details then only provide requested information.
            Provide the analysis in a structured format as outlined above.
            Please focus solely on extracting the requested information.
            Ensure accuracy in counting the 'Total number of resumes','Number of Resumes suited for Data Science job profile','Number of Not suitable resumes'.

            Here is the provided data: {data}

            """

    response_text = model.generate_content(prompt3, generation_config=genai.types.GenerationConfig(
        candidate_count=1,
        temperature=0)
                                           )

    return response_text.text

#This functionality is to be trigerred from frontend for downloading csv.
def download_button(object_to_download, download_filename, button_text='Download as CSV'):
    csv = object_to_download.to_csv(index=False)

    b64 = base64.b64encode(csv.encode()).decode()

    href = f'<a href="data:file/csv;base64,{b64}" download="{download_filename}">{button_text}</a>'

    return href


def download_button_excel(object_to_download, download_filename, button_text='Download'):
    object_to_download.seek(0)
    timestamp = datetime.now().strftime('%Y/%m/%d/%H:%M:%S')
    download_filename = f'resume_summaries_{timestamp}.xlsx'
    b64 = base64.b64encode(object_to_download.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{download_filename}">{button_text}</a>'

#Function to process the firstbox
def prathambox(firstbox):
    x = firstbox
    return x

#Function to process the secondbox
def dutiyabox(skills_to_search):
    y = skills_to_search
    return y
#Function to process the thirdbox
def tritiyabox(folder_path):
    global resume_folder
    if folder_path:
        resume_folder = rf"{folder_path}"
    return resume_folder
#function to split the optionalskills
def fourthbox(skills_to_search):
    global Optional_skills
    Optional_skills = [skill.strip() for skill in skills_to_search.split(",")]

    return Optional_skills

#This function shows userInput summary.
def userInput_summary(firstbox, skills_to_search,folder_path):
    if folder_path:
        resume_folder = rf"{folder_path}"


    resumes_dict = load_resumes_as_dict(resume_folder)

    # Calculate number of files in folder
    num_files = len(resumes_dict)
    # Display number of files to user
    xy = firstbox + ", " + skills_to_search

    library = []
    library.append([num_files, xy, folder_path])
    df_library = pd.DataFrame(library, columns=["# Resumes to Screen", "Required Skillset to Filter",
                                                "Path of the Repository of the Resumes"])

    return df_library.to_json()

#Function to return df_summaries
def dataFr(Optional_skills,resume_folder):

    datascience_skills = ['Statistical analysis', 'Feature engineering', 'Model development', 'Model validation',
                          'Model deployment', 'Predictive modeling', 'Time series analysis',
                          'Natural language processing', 'Big data technologies', 'Quantitative research',
                          'Business intelligence', 'Predictive analytics', 'Descriptive analytics',
                          'Prescriptive analytics', 'Sentiment analysis', 'Text mining', 'Anomaly detection',
                          'Classification', 'Regression', 'Clustering', 'Dimensionality reduction',
                          'Ensemble methods', 'Neural networks', 'Convolutional Neural Networks (CNNs)',
                          'Recurrent Neural Networks (RNNs)', 'Transfer learning', 'Reinforcement learning',
                          'Optimization algorithms', 'Hyperparameter tuning', 'Cross-validation',
                          'Statistical modeling', 'Bayesian inference', 'Markov models', 'Social media analysis']

    batch_size = 10
    files = os.listdir(resume_folder)
    num_filess = len(files)
    num_batches = (num_filess + batch_size - 1) // batch_size

    for i in range(num_batches):
        start_i = i * batch_size
        end_i = min((i + 1) * batch_size, num_filess)
        batch_files = files[start_i:end_i]
        resumes_dict = load_resumes_as_dict2(resume_folder, batch_files)
        # st.write(resumes_dict)
        summary_data = generate_resume_summary(resumes_dict, Optional_skills, datascience_skills,resume_folder)

    df_summaries = pd.DataFrame(summary_data,

                                columns=['File Name', 'Name of the candidate', 'Summary', 'Relevant Experience', 'Skills',
                                         'GenAI','Data Science', 'Cloud', 'Github',
                                         'Programming Languages', 'Domain', 'Recommendation',
                                         'Additional Skills', 'Alternate Recommendation', 'View PDF','View Docx'])

    df_summaries['Additional Skills'] = ''
    df_summaries['Alternate Recommendation'] = 'No'  # Default to 'No'

    for i in Optional_skills:
        mask = df_summaries['Skills'].str.contains(i, case=False)
        recommendation_mask = df_summaries['Recommendation'].str.startswith('Yes', na=False)

        # Check if the skill is not already present in 'Additional Skills'
        not_present_mask = ~df_summaries['Additional Skills'].str.contains(i, case=False) | df_summaries[
            'Additional Skills'].isna()

        # Update 'Alternate Recommendation' based on conditions
        df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Alternate Recommendation'] = 'Yes'

        # Append the skill only if it's not already present
        df_summaries.loc[mask & recommendation_mask & not_present_mask, 'Additional Skills'] = df_summaries.loc[
            mask & recommendation_mask & not_present_mask, 'Additional Skills'].apply(
            lambda x: x + ',' + i if pd.notna(x) and x != '' else i)

    # Update 'Recommendation' to 'No' if 'Additional Skills' is empty
    df_summaries.loc[df_summaries['Additional Skills'].eq('') | df_summaries[
        'Additional Skills'].isna(), 'Alternate Recommendation'] = 'No'

    return  df_summaries

#function to return genai_count
def genAICount(noob):
    genai_count = noob['Recommendation'].str.contains('Recommended for GenAI profile').sum()

    return genai_count

# function to return datscience_count
def dsCount(noob):
    data_science_count = noob['Recommendation'].str.contains(
        'Recommended for Data Science profile').sum()

    return data_science_count

# function to return preffered_count
def altrCount(noob):
    alternate_count = noob['Alternate Recommendation'].str.contains('Yes').sum()

    return  alternate_count

#function to return borderline_count
def bdrCount(noob):
    borderline_count = noob['Recommendation'].str.contains('Borderline').sum()

    return  borderline_count

#function to return resume summary
def resuSumm(noob, skill, additionalskill, path):
    summary = []
    total_resumes = len(noob)
    recommended_count = noob['Recommendation'].str.contains('Yes').sum()
    borderline_count = noob['Recommendation'].str.contains('Borderline').sum()
    notrecommended_count = noob['Recommendation'].str.contains('No').sum()

    data_count = {"Total": str(total_resumes), "Recommended": str(recommended_count),
                  "Borderline": str(borderline_count), "Not Recommended": str(notrecommended_count)}
    data_count = {"Resume Classification": data_count}
    json_data = json.dumps(data_count)


    genai_count = noob['Recommendation'].str.contains('Yes. Recommended for GenAI profile.').sum()
    genai_borderline_count = noob['Recommendation'].str.contains('Borderline. Recommended for GenAI profile.').sum()
    filtered_skills_genai = noob['GenAI'].dropna().replace('None', '').replace('Not Mentioned','').replace('Error in processing', '').replace('Could not read', '')
    all_genai_skills = filtered_skills_genai.str.split(',').explode().str.strip().str.rstrip(';')
    all_genai_skills = all_genai_skills[all_genai_skills != '']
    trend_genai_df = pd.DataFrame({'skills': all_genai_skills})
    top5_ganai_skills = trend_genai_df['skills'].value_counts().nlargest(5).index
    top5_ganai_skills_dict = top5_ganai_skills.tolist()
    top5_ganai_skills_dict = ', '.join(top5_ganai_skills_dict)

    genai_ui = {"Recommended": str(genai_count), "Borderline": str(genai_borderline_count),
                "Trend": top5_ganai_skills_dict}
    genai_ui = {"GenAI count": genai_ui}
    genai_ui_json = json.dumps(genai_ui)


    data_science_count = noob['Recommendation'].str.contains('Yes. Recommended for Data Science profile.').sum()
    datascience_borderline_count = noob['Recommendation'].str.contains('Borderline. Recommended for Data Science profile.').sum()
    filtered_skills_ds = noob['Data Science'].dropna().replace('None', '').replace('Not Mentioned','').replace('Error in processing', '').replace('Could not read', '')
    all_datascience_skills = filtered_skills_ds.str.split(',').explode().str.strip().str.rstrip(';')
    all_datascience_skills = all_datascience_skills[all_datascience_skills != '']
    trend_datascience_df = pd.DataFrame({'skills': all_datascience_skills})
    top5_datasceince_skills = trend_datascience_df['skills'].value_counts().nlargest(5).index
    top5_datasceince_skills_dict = top5_datasceince_skills.tolist()
    top5_datasceince_skills_dict = ', '.join(top5_datasceince_skills_dict)

    ds_ui = {"Recommended": str(data_science_count), "Borderline": str(datascience_borderline_count),
             "Trend": top5_datasceince_skills_dict}
    ds_ui = {"Data Science count": ds_ui}
    ds_ui_json = json.dumps(ds_ui)


    alternate_count = noob['Alternate Recommendation'].str.contains('Yes').sum()
    genai_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    genai_additional_count = genai_additional_count[
        genai_additional_count['Recommendation'].str.contains('Recommended for GenAI profile')]
    genai_additional_count = len(genai_additional_count)
    ds_additional_count = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    ds_additional_count = ds_additional_count[
        ds_additional_count['Recommendation'].str.contains('Recommended for Data Science profile.')]
    ds_additional_count = len(ds_additional_count)

    additional_ui = {"Recommended": str(alternate_count), "GenAI Recommended": str(genai_additional_count),
                     "Data Science Recommended": str(ds_additional_count)}
    additional_ui = {"Additional count": additional_ui}
    additional_ui_json = json.dumps(additional_ui)

    table_ui = json.dumps({"serial":str(total_resumes),"skillset":str(skill),"additionalSkills":str(additionalskill),"path":str(path)})

    # summary.append([total_resumes, m1, m2, m3])
    #
    # df_new = pd.DataFrame(summary, columns=['Total Number of Resumes', 'GenAI', 'Data Science',
    #                                         'Profile screened for Additional Skills'])

    return json_data,genai_ui_json,ds_ui_json,additional_ui_json,table_ui

#function for GenAI resume results
def genAIres(noob, m1):
    #global gensa
    genai_df = noob[noob['Recommendation'].str.contains('Recommended for GenAI profile')]
    show_genai_df = genai_df[
        ['Name of the candidate', 'Skills','GenAI', 'Data Science', 'Cloud','Programming Languages','View PDF','Recommendation','File Name']]
    if m1 == 0:
        print("No GenAI related resumes found !")
    else:
        show_genai_df
    return show_genai_df.to_json(orient ='records')

#function for dataScience resume results
def DSres(noob, m1):
    #global dssa
    data_science_df = noob[
        noob['Recommendation'].str.contains('Recommended for Data Science profile')]
    show_data_science_df = data_science_df[
        ['Name of the candidate', 'Skills','GenAI', 'Data Science', 'Cloud','Programming Languages','View PDF','Recommendation','File Name']]
    if m1 == 0:
        print("No GenAI related resumes found !")
    else:
        show_data_science_df
    return show_data_science_df.to_json(orient ='records')

#function for alternate resume results
def altrres(noob, m1):
    global altsa
    alternate_df = noob[noob['Alternate Recommendation'].str.contains('Yes')]
    show_alternate_df = alternate_df[
        ['Name of the candidate', 'Skills','GenAI', 'Data Science', 'Cloud','Programming Languages','View PDF']]
    if m1 == 0:
        print("No GenAI related resumes found !")
    else:
        show_alternate_df
    return show_alternate_df.to_json(orient ='records')

# def main():
#
#
#             genai_df = df_summaries[df_summaries['Recommendation'].str.contains('Recommended for GenAI profile')]
#             show_genai_df = genai_df[
#                 ['Name of the candidate', 'GenAI Related', 'Data Science Related', 'Cloud Related']]
#
#             data_science_df = df_summaries[
#                 df_summaries['Recommendation'].str.contains('Recommended for Data Science profile')]
#             show_data_science_df = data_science_df[
#                 ['Name of the candidate', 'GenAI Related', 'Data Science Related', 'Cloud Related']]
#
#             alternate_df = df_summaries[df_summaries['Alternate Recommendation'].str.contains('Yes')]
#             show_alternate_df = alternate_df[
#                 ['Name of the candidate', 'GenAI Related', 'Data Science Related', 'Cloud Related']]
#
#             borderline_df = df_summaries[df_summaries['Recommendation'].str.contains('Borderline')]
#             show_borderline_df = borderline_df[
#                 ['Name of the candidate', 'GenAI Related', 'Data Science Related', 'Cloud Related']]
#
#             html1 = df_new.to_html(index=False)
#
#             # Use Streamlit's markdown to display the table
#             # st.markdown(html1, unsafe_allow_html=True)
#             st.markdown(f"<style>.dataframe{{width:100%;}}</style>{html1}", unsafe_allow_html=True)
#             st.markdown(" ")
#
#             df_recommended = df_summaries[df_summaries['Recommendation'].str.contains('Recommended')]
#
#             recommended_resumes = {}
#             for index, row in df_recommended.iterrows():
#                 file_name = row['File Name']
#                 if file_name in resumes_dict:
#                     _, content = resumes_dict[file_name]
#                     recommended_resumes[file_name] = content
#
#             # st.write(recommended_resumes)
#             # questions = generate_questions(recommended_resumes)
#
#             # df_questions = pd.DataFrame(questions, columns=['File Name','Questions'])
#
#             # st.table(df_new)
#             with st.expander("GenAI"):
#                 # show_df1 = show_genai_df.style.apply(
#                 #     lambda row: [
#                 #         'background-color: #00FF00' if 'Yes' in val else 'background-color: #FF0000' if 'No' in val else ''
#                 #         for val in row],
#                 #     axis=1,
#                 #     subset=['Recommendation']
#                 # )
#                 if genai_count == 0:
#                     st.write("No GenAI related resumes found !")
#                 else:
#                     st.table(show_genai_df)
#
#             with st.expander("Data Science"):
#                 # show_df2 = show_data_science_df.style.apply(
#                 #     lambda row: [
#                 #         'background-color: #00FF00' if 'Yes' in val else 'background-color: #FF0000' if 'No' in val else ''
#                 #         for val in row],
#                 #     axis=1,
#                 #     subset=['Recommendation']
#                 # )
#                 if data_science_count == 0:
#                     st.write("No Data Science related resumes found !")
#                 else:
#                     st.table(show_data_science_df)
#
#             with st.expander("Profile screened for additional skills"):
#                 if alternate_count == 0:
#                     st.write("No resumes found !")
#                 else:
#                     st.table(show_alternate_df)
#
#
#
#
# if __name__ == "__main__":
#     main()



                

            
